from docx import Document
from docx.shared import Pt
from inflection import *
from num_to_word import *
from globals import *


class doc():
    
    def __init__(self, data, allowances):
        self.data = data
        self.allow = allowances
        
        self.contract_info = contract_info
        self.allowances = allowances
        

    def updateAllow(self, data):
        pass
    
    def fill_allow(self):
        pass
    
    def fill_contract(self):
        document = Document("docs/Contract.docx")
        style = document.styles['Normal']
        font = style.font
        font.name = 'Stymie Lt BT'
        font.size = Pt(11)
        
        date = self.contract_info["date"]
        client = self.contract_info["client"]
        est_num = self.contract_info["estimate_num"]
        start = self.contract_info["start_date"]
        end = self.contract_info["end_date"]
        payment_amount = num_to_words(self.contract_info["estiamte_amount"])
        hourly_rate = self.contract_info["hourly_rate"]
        time = 0;
        if self.contract_info["time_important"] > 0:
            time = 1
        deposit = self.contract_info["deposit"]
 
        print(document.paragraphs[4].text)     
        
        #----------------------------------------------------------------------#  
        document.paragraphs[2].clear()
        document.paragraphs[2].add_run(text = date, style = None)
        
        document.paragraphs[4].clear()
        document.paragraphs[4].add_run(text = client, style = None)
        #----------------------------------------------------------------------#
        est_num_text = f"work outlined in Estimate #{est_num} detailed below in this contract."
        document.paragraphs[7].clear()
        document.paragraphs[7].add_run(text = est_num_text, style = None)    
        #----------------------------------------------------------------------#
        start_date_text = f"Anticipated start date: {start}"
        end_date_text = f"Anticipated completion date: {end}"
        document.paragraphs[24].clear()
        document.paragraphs[24].add_run(text = start_date_text, style = None)
        
        document.paragraphs[25].clear()
        document.paragraphs[25].add_run(text = end_date_text, style = None)
        #----------------------------------------------------------------------#
        deposit_text = f"Deposit: {deposit}"
        document.paragraphs[7].clear()
        document.paragraphs[7].add_run(text = est_num_text, style = None)   
        #----------------------------------------------------------------------#
        payment_paragraph_1 = f'''All of the above work is to be completed in a substantial and workmanlike manner for the sum of'''

        payment_paragraph_2 = f' {payment_amount} '
        
        payment_paragraph_3 = f'''
There will be minor alterations to the contract sum based on changes during construction. 
Any alterations or deviation of the work described and as above will be executed upon written 
order for the same and will be added or deducted from the sum quoted in this contract or will be 
performed on a time and material basis at ${hourly_rate}/man/hour.'''

        document.paragraphs[27].clear()
        document.paragraphs[27].add_run(text = payment_paragraph_1, style = None)
        document.paragraphs[27].add_run(text = payment_paragraph_2, style = None).bold = True
        document.paragraphs[27].add_run(text = payment_paragraph_3, style = None)
        
        #----------------------------------------------------------------------#
        
        '''
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        picture = run.add_picture("image.jpg")
        '''
             
        document.save("docs/Contract.docx")

